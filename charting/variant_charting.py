import pandas as pd
import numpy as np
import sort
import display
import docx

cond = [
    ['breast', 'asian_census', 'census'],
    ['hepatocellular_carcinoma', 'asian_census', 'census'],
    ['lung', 'asian_census', 'census'],
    ['rectum', 'asian_census', 'census'],
    ['thyroid', 'asian_census', 'census']
]

GENE_MAX_NUM = 50

def data_process():
    for j in range(0,5):
        for i in range (1,3):
            df = pd.read_csv(f"raw_output\census_mutation_output\cosmic_{cond[j][0]}_cancer_{cond[j][i]}_mutation.csv")
            df = df[df['Tier'] == 1]    
            df = sort._sort_by_variant(df)
            df.to_csv(f"output\{cond[j][0]}_cancer\csv\(variant)sorted_cosmic_{cond[j][0]}_cancer_{cond[j][i]}_mutation.csv", index=None, header=True)

def write_to_doc(asia_gene, world_gene, cancer_name):
    mydoc = docx.Document()
    mydoc.add_paragraph(f"Top AA variants of {cancer_name} cancer are (n={GENE_MAX_NUM}):")
    para = "Asia:\n"
    data = pd.DataFrame(
        {'list' : asia_gene['list'],
         'ratio' : asia_gene['ratio']
        }).head(min(len(asia_gene['list']), GENE_MAX_NUM))
    data = data.sort_values('ratio', ascending=False).reset_index(drop=True)
    for k in range(0, len(data)):
        para += f"{data['list'][k]} ({round(data['ratio'][k]*100,4)}%)\n"
    para += "\n\n"
    para += "The world:\n"
    data = data[0:0]
    data['list'] = world_gene['list']
    data['ratio'] = world_gene['ratio']
    for k in range(0, len(data)):
        para += f"{data['list'][k]} ({round(data['ratio'][k]*100,4)}%)\n"
    mydoc.add_paragraph(para)
    mydoc.save(f"output\{cancer_name}_cancer\doc\(variant){cancer_name}_cancer_list.docx")



def frequency_cal():
    for j in range(0,5):
        ratio = dict()
        ratio_census = dict()
        for i in range (1,3):
            df = pd.read_csv(f"output\{cond[j][0]}_cancer\csv\(variant)sorted_cosmic_{cond[j][0]}_cancer_{cond[j][i]}_mutation.csv")
            df = df.reindex(['ID_tumour', 'Mutation AA'], axis=1)[df['Mutation AA'] != "p.?"]
            list = df['Mutation AA'].unique() #a list of unique AA mutations in the data
            freq_list = [0]*len(list)
            ratio_list = [1.0]*len(list)
            aa_ratio = {
                        'list' : list, 
                        'freq': freq_list,  #actual number of occurances for each gene with unique id_tumour values
                        'ratio': ratio_list, #ratio of each gene among the sum
                        }
            sum_aa = 0 #number of all genes with unique id_tumour values 
            for k in range(0, len(list)):
                aa_ratio['freq'][k] = sort._count_real_cds_freq(list[k],df)
                sum_aa += aa_ratio['freq'][k]   
            for k in range(0, len(list)):
                aa_ratio['ratio'][k] = aa_ratio['freq'][k] / sum_aa
            if i==1:
                ratio = aa_ratio
            else:
                ratio_census = aa_ratio
        display._chart_create(ratio, ratio_census, cond[j][0], 'v', 'Asia', 'The World')
        #write_to_doc(ratio, ratio_census, cond[j][0])

#data_process()
frequency_cal()
            
            






